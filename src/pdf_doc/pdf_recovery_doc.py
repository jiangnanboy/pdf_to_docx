import os
import cv2
import numpy as np
import time
from docxcompose.composer import Composer
from docx import Document
import docx
# from win32com import client
from paddleocr.ppstructure.recovery.recovery_to_doc import sorted_layout_boxes, convert_info_docx

from paddleocr.ppstructure.predict_system import save_structure_res

from paddleocr import PPStructure

from src.util.image_process import pdf_image

def pdf2doc(pdf_file):
    '''
    pdf2doc
    :param pdf_file:
    :return:
    '''
    table_engine = PPStructure(recovery=True, lang='ch', enable_mkldnn=True)
    start_time = time.time()
    if pdf_file.endswith('.pdf') or pdf_file.endswith('.PDF'):
        pdf_file_base_name = os.path.splitext(pdf_file)[0]
        image_path = pdf_file_base_name + '/' + 'images'

        # 1.pdf to image
        pdf_image(pdf_file, image_path)
        image_files = os.listdir(image_path)

        # 2.image to docx
        for img_file in image_files:
            if img_file.endswith('png'):
                img = cv2.imdecode(np.fromfile(os.path.join(image_path, img_file), dtype=np.uint8), -1)
                img = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
                if img is not None:
                    result = table_engine(img)
                    save_structure_res(result, pdf_file_base_name, img_file.split('.')[0])
                    h, w, _ = img.shape
                    res = sorted_layout_boxes(result, w)
                    convert_info_docx(img, res, pdf_file_base_name, img_file.split('.')[0])

        # 3.merge all docx
        merge_docx_v1(pdf_file_base_name)

    end_time = time.time()
    print('pdf转换doc时间: {}'.format(end_time - start_time))

def merge_docx_v1(docx_dir):
    '''
    merge all docx
    :param docx_dir:
    :return:
    '''
    doc_base_name = os.path.basename(docx_dir)
    docx_path = os.listdir(docx_dir)
    docx_files_list = [docx_file for docx_file in docx_path if docx_file.endswith('docx')]
    docx_files_list = [os.path.join(docx_dir, docx_file) for docx_file in docx_files_list]
    docx_files_list = sorted(docx_files_list)
    frist_docx = docx_files_list[0]
    docx_file_count = len(docx_files_list)
    docx_master = Document(frist_docx)
    composer = Composer(docx_master)
    for i in range(1, docx_file_count):
        doc_temp = Document(docx_files_list[i])
        doc_temp.add_page_break()
        composer.append(doc_temp)
    # composer.save(os.path.join(docx_dir, doc_base_name + ".docx"))
    docx_master.save(os.path.join(docx_dir, doc_base_name + '.docx'))

def merge_docx_v2(docx_dir):
    '''
    merge all docx
    :param docx_dir:
    :return:
    '''
    docx_path = os.listdir(docx_dir)
    docx_files_list = [docx_file for docx_file in docx_path if docx_file.endswith('docx') ]
    docx_files_list = [os.path.join(docx_dir, docx_file) for docx_file in docx_files_list]
    docx_files_list = sorted(docx_files_list)

    merge_doc = docx.Document()
    for docx_file in docx_files_list:
        docx_doc = docx.Document(docx_file)
        for word_body in docx_doc.element.body:
            merge_doc.element.body.append(word_body)
    merge_doc.save(os.path.join(docx_dir, os.path.basename(docx_dir) + '.docx'))

def merge_docx_v3(docx_dir):
    '''
    merge all docx
    :param docx_dir:
    :return:
    '''
    # docx_path = os.listdir(docx_dir)
    # docx_files_list = [docx_file for docx_file in docx_path if docx_file.endswith('docx')]
    # docx_files_list = [os.path.join(docx_dir, docx_file) for docx_file in docx_files_list]
    # docx_files_list = sorted(docx_files_list)
    #
    # word = client.gencache.EnsureDispatch('Word.Application')
    # word.Visible = True
    # new_document = word.Documents.Add()
    # for fn in docx_files_list[::-1]:
    #     new_document.Application.Selection.Range.InsertFile(fn)
    # new_document.SaveAs(os.path.join(docx_dir, os.path.basename(docx_dir) + '.docx'))
    # new_document.Close()
    # word.Quit()


if __name__ == '__main__':
    file_path = os.path.abspath(os.path.join(os.getcwd(), "../.."))
    data_path = os.path.join(file_path, 'data')
    pdf_path = os.path.join(data_path, '中文文本自动校对综述.pdf')
    pdf2doc(pdf_path)

